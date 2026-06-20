"""
Phase 4 转换管道 — 快速验证脚本

用法:
    python tests/test_phase4_quick.py
    
功能:
1. 创建测试用 DOCX 文件（含标题、表格）
2. 运行转换管道
3. 输出结果摘要
"""

import os
import sys
from pathlib import Path

# 确保项目根目录在路径中
sys.path.insert(0, str(Path(__file__).parent.parent))

try:
    from docx import Document
except ImportError:
    print("⚠️ python-docx 未安装，运行 'pip install python-docx'")
    sys.exit(1)

from src.converters.docx_converter import DocxConverter
from src.conversion_pipeline import ConversionPipeline


def create_test_docx():
    """
    创建测试用 DOCX 文件。
    
    Returns:
        测试文件路径
    """
    test_dir = Path(__file__).parent / "test_files"
    os.makedirs(test_dir, exist_ok=True)
    
    filepath = test_dir / "test_document.docx"
    doc = Document()
    
    # 添加标题
    doc.add_heading("测试文档", level=0)
    doc.add_paragraph("这是一个用于验证 Phase 4 转换管道的测试文档。")
    
    # 第一章
    doc.add_heading("第一章：概述", level=1)
    doc.add_paragraph("这是第一章的内容，介绍项目背景和目标。")
    doc.add_paragraph("**加粗文本**和*斜体文本*的格式验证。")
    
    # 第二章
    doc.add_heading("第二章：市场分析", level=1)
    doc.add_paragraph("以下是市场数据表格：")
    
    # 添加表格
    table = doc.add_table(rows=4, cols=3, style='Table Grid')
    headers = ['指标', '2024年', '2025年']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    
    data = [
        ['营收 (万元)', '1000', '1500'],
        ['利润 (万元)', '200', '400'],
        ['增长率', '10%', '50%'],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_text in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = cell_text
    
    # 第二章子章节
    doc.add_heading("2.1 行业趋势", level=2)
    doc.add_paragraph("AI 芯片市场规模持续扩大，预计 2026 年达到 XXX 亿元。")
    
    doc.add_heading("2.2 竞品分析", level=2)
    doc.add_paragraph("主要竞争对手包括 A 公司、B 公司和 C 公司。")
    
    # 第三章
    doc.add_heading("第三章：产品路线图", level=1)
    doc.add_paragraph("Q1: 完成 MVP 开发")
    doc.add_paragraph("Q2: 用户测试与迭代")
    doc.add_paragraph("Q3: 正式发布")
    
    # 保存
    doc.save(str(filepath))
    print(f"✅ 创建测试文件: {filepath}")
    return filepath


def test_docx_converter():
    """
    测试 DOCX 转换器。
    """
    print("\n" + "=" * 60)
    print("🧪 测试 1: DOCX 转换器")
    print("=" * 60)
    
    # 创建测试文件
    test_file = create_test_docx()
    
    # 运行转换
    converter = DocxConverter(max_heading_level=3)
    result = converter.convert(str(test_file))
    
    # 输出结果摘要
    print(f"\n📊 转换结果:")
    print(f"   - 章节数: {len(result['sections'])}")
    print(f"   - 标题数: {len(result['headings'])}")
    print(f"   - 表格数: {len(result['tables'])}")
    print(f"   - 估算页数: {result['total_pages']}")
    
    # 打印章节列表
    print("\n📑 章节结构:")
    for section in result['sections']:
        indent = "  " * (section.heading_level - 1)
        print(f"{indent}- [{section.heading_level}] {section.title}")
    
    # 打印表格内容
    if result['tables']:
        print("\n📋 提取的表格:")
        for i, table in enumerate(result['tables']):
            print(f"   表格 {i + 1}:")
            print(f"      表头: {' | '.join(table.headers)}")
            for row in table.rows[:3]:  # 最多显示 3 行
                print(f"      数据: {' | '.join(row)}")
    
    return True


def test_conversion_pipeline():
    """
    测试转换管道。
    """
    print("\n" + "=" * 60)
    print("🧪 测试 2: 转换管道")
    print("=" * 60)
    
    # 创建测试文件
    test_file = create_test_docx()
    
    # 运行管道（不生成树状 MD，只输出到控制台）
    pipeline = ConversionPipeline(
        db_path=str(Path(__file__).parent.parent / "wiki_db"),
        output_dir=None  # 不生成文件
    )
    
    result = pipeline.convert(str(test_file))
    
    print(f"\n📊 管道结果:")
    print(f"   - 文件格式: {result.conversion_method.upper()}")
    print(f"   - 章节数: {len(result.sections)}")
    print(f"   - 质量评分: {result.quality_score:.2f}")
    print(f"   - LLM 修复需求: {'是' if result.metadata.get('needs_llm_fix') else '否'}")
    
    return True


def test_tree_generation():
    """
    测试树状 MD 生成。
    """
    print("\n" + "=" * 60)
    print("🧪 测试 3: 树状 MD 生成")
    print("=" * 60)
    
    # 创建测试文件
    test_file = create_test_docx()
    
    # 运行管道（生成树状 MD）
    output_dir = str(Path(__file__).parent / "test_files" / "output")
    pipeline = ConversionPipeline(
        db_path=str(Path(__file__).parent.parent / "wiki_db"),
        output_dir=output_dir
    )
    
    result = pipeline.convert(str(test_file))
    
    print(f"\n📁 生成的文件:")
    for md_file in result.md_files_created:
        rel_path = os.path.relpath(md_file, output_dir)
        print(f"   - {rel_path}")
    
    # 显示 _index.md 内容
    index_path = Path(output_dir) / "test_document" / "_index.md"
    if index_path.exists():
        print("\n📄 _index.md 预览:")
        with open(index_path, 'r', encoding='utf-8') as f:
            content = f.read()
            # 只显示前 30 行
            lines = content.split('\n')[:30]
            print('\n'.join(lines))
    
    return True


if __name__ == "__main__":
    print("🚀 Phase 4 转换管道 — 快速验证")
    print("=" * 60)
    
    try:
        # 测试 1: DOCX 转换器
        test_docx_converter()
        
        # 测试 2: 转换管道
        test_conversion_pipeline()
        
        # 测试 3: 树状 MD 生成
        test_tree_generation()
        
        print("\n" + "=" * 60)
        print("✅ 所有测试通过！")
        print("=" * 60)
    except Exception as e:
        print(f"\n❌ 测试失败: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
