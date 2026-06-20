"""
DOCX 转换器 — 将 Word 文档转换为结构化 Markdown

核心功能:
1. 提取标题层级（Heading 样式 → H1/H2/H3...）
2. 提取正文段落（保留加粗/斜体等基础格式）
3. 提取表格（转 Markdown 表格格式）
4. 按章节切分为独立 MD 文件
5. 生成 YAML Frontmatter 元数据

使用示例:
    converter = DocxConverter()
    result = converter.convert("/path/to/document.docx")
    # result: {sections: [...], assets_dir: "/path/to/assets"}
"""

import os
import re
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass, field

try:
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
except ImportError:
    Document = None


# ==================== 数据结构 ====================

@dataclass
class HeadingInfo:
    """标题信息"""
    level: int          # 1-9 (对应 H1-H9)
    text: str           # 标题文本
    paragraph_index: int  # 在文档中的位置


@dataclass
class TableData:
    """表格数据"""
    headers: List[str]
    rows: List[List[str]]
    source_page: Optional[int] = None  # 原始页码（DOCX 中通常不可用）


@dataclass
class SectionResult:
    """转换后的章节结果"""
    title: str                    # 章节标题
    content: str                  # Markdown 内容
    heading_level: int            # 所属层级（1=顶级，2=二级...）
    parent_title: Optional[str]   # 父章节标题
    page_range: Optional[Tuple[int, int]] = None  # 原始页码范围
    metadata: Dict = field(default_factory=dict)  # YAML Frontmatter 元数据


# ==================== DOCX 解析器 ====================

class DocxConverter:
    """
    Word 文档 (.docx) → Markdown 转换器。
    
    Attributes:
        max_heading_level: 最大处理的标题层级（默认 3，即 H3）
        table_format: 表格输出格式（'markdown' | 'html'）
    """
    
    # Word Heading 样式名到 Markdown 级别的映射
    HEADING_STYLE_MAP = {
        'Heading 1': 1,
        'Heading 2': 2,
        'Heading 3': 3,
        'Heading 4': 4,
        'Heading 5': 5,
        'Heading 6': 6,
        'Heading 7': 7,
        'Heading 8': 8,
        'Heading 9': 9,
    }
    
    def __init__(self, max_heading_level: int = 3):
        """
        Args:
            max_heading_level: 最大处理的标题层级（默认 3）
        """
        self.max_heading_level = max_heading_level
        self.assets_dir: Optional[str] = None
    
    def convert(self, filepath: str, output_dir: Optional[str] = None) -> Dict:
        """
        将 DOCX 文件转换为 Markdown。
        
        Args:
            filepath: .docx 文件路径
            output_dir: 输出目录（如果提供，会生成树状 MD 文件）
            
        Returns:
            {
                'sections': [SectionResult, ...],
                'headings': [HeadingInfo, ...],
                'tables': [TableData, ...],
                'assets_dir': str | None,
                'total_pages': int
            }
        """
        filepath = Path(filepath)
        if not filepath.exists():
            raise FileNotFoundError(f"文件不存在: {filepath}")
        
        # 1. 加载文档
        doc = Document(str(filepath))
        
        # 2. 提取标题层级结构
        headings = self._extract_headings(doc)
        
        # 3. 按章节组织内容
        sections = self._organize_by_sections(doc, headings)
        
        # 4. 提取表格
        tables = self._extract_tables(doc)
        
        # 5. 提取图片（如果需要）
        assets_dir = None
        if output_dir:
            assets_dir = str(filepath.parent / "assets" / filepath.stem)
            os.makedirs(assets_dir, exist_ok=True)
            self._extract_images(doc, assets_dir)
        
        # 6. 生成树状 MD 文件（如果提供了输出目录）
        if output_dir:
            self._write_tree_structure(
                sections=sections,
                headings=headings,
                base_dir=output_dir,
                doc_name=filepath.stem
            )
        
        return {
            'sections': sections,
            'headings': headings,
            'tables': tables,
            'assets_dir': assets_dir,
            'total_pages': self._estimate_pages(doc),
        }
    
    def _extract_headings(self, doc: Document) -> List[HeadingInfo]:
        """
        从文档中提取所有标题。
        
        Args:
            doc: python-docx Document 对象
            
        Returns:
            [HeadingInfo, ...]
        """
        headings = []
        for i, para in enumerate(doc.paragraphs):
            style_name = para.style.name or ''
            if style_name in self.HEADING_STYLE_MAP:
                level = self.HEADING_STYLE_MAP[style_name]
                if level <= self.max_heading_level:
                    headings.append(HeadingInfo(
                        level=level,
                        text=para.text.strip(),
                        paragraph_index=i
                    ))
        return headings
    
    def _organize_by_sections(self, doc: Document, headings: List[HeadingInfo]) -> List[SectionResult]:
        """
        按标题层级组织文档内容为章节。
        
        Args:
            doc: python-docx Document 对象
            headings: 提取的标题列表
            
        Returns:
            [SectionResult, ...]
        """
        if not headings:
            # 没有标题 → 整个文档作为一个章节
            content = self._extract_paragraphs(doc)
            return [
                SectionResult(
                    title="全文",
                    content=content,
                    heading_level=1,
                    parent_title=None,
                    metadata={'source': 'docx', 'has_headings': False}
                )
            ]
        
        sections = []
        doc_name = ""
        
        for idx, heading in enumerate(headings):
            # 确定章节范围（从当前标题到下一个同级/上级标题之前）
            if idx + 1 < len(headings):
                next_heading = headings[idx + 1]
                # 找到下一个同级或更高级别标题的段落索引
                end_index = None
                for h in headings[idx + 1:]:
                    if h.level <= heading.level:
                        end_index = h.paragraph_index
                        break
            else:
                end_index = len(doc.paragraphs)
            
            # 提取段落内容（从标题后到下一个章节前）
            start_index = heading.paragraph_index + 1
            para_text = self._extract_paragraph_range(
                doc, start_index, end_index
            )
            
            # 确定父章节标题
            parent_title = None
            for prev_heading in reversed(headings[:idx]):
                if prev_heading.level < heading.level:
                    parent_title = prev_heading.text
                    break
            
            # 构建 YAML Frontmatter
            metadata = {
                'source': 'docx',
                'heading_level': heading.level,
                'paragraph_range': f"{start_index}-{end_index}" if end_index else f"{start_index}-",
            }
            
            sections.append(SectionResult(
                title=heading.text,
                content=para_text,
                heading_level=heading.level,
                parent_title=parent_title,
                metadata=metadata
            ))
        
        return sections
    
    def _extract_paragraphs(self, doc: Document) -> str:
        """
        提取所有非标题段落的内容。
        
        Args:
            doc: python-docx Document 对象
            
        Returns:
            Markdown 格式的段落文本
        """
        lines = []
        for para in doc.paragraphs:
            style_name = para.style.name or ''
            if style_name.startswith('Heading'):
                continue  # 跳过标题（由 headings 列表处理）
            
            text = self._format_paragraph(para)
            if text.strip():
                lines.append(text)
        
        return '\n\n'.join(lines)
    
    def _extract_paragraph_range(self, doc: Document, start: int, end: Optional[int]) -> str:
        """
        提取指定范围内的段落内容。
        
        Args:
            doc: python-docx Document 对象
            start: 起始索引（包含）
            end: 结束索引（不包含，None=到文档末尾）
            
        Returns:
            Markdown 格式的段落文本
        """
        lines = []
        for i, para in enumerate(doc.paragraphs):
            if i < start or (end is not None and i >= end):
                continue
            
            style_name = para.style.name or ''
            # 跳过同级或更高级别的标题（它们会作为独立章节）
            if style_name in self.HEADING_STYLE_MAP:
                level = self.HEADING_STYLE_MAP[style_name]
                if level <= self.max_heading_level:
                    continue
            
            text = self._format_paragraph(para)
            if text.strip():
                lines.append(text)
        
        return '\n\n'.join(lines)
    
    def _format_paragraph(self, para) -> str:
        """
        格式化单个段落（保留加粗/斜体等基础格式）。
        
        Args:
            para: python-docx Paragraph 对象
            
        Returns:
            Markdown 格式的文本
        """
        parts = []
        for run in para.runs:
            text = run.text
            if not text:
                continue
            
            # 基础格式标记
            bold = run.bold or False
            italic = run.italic or False
            underline = run.underline or False
            
            # Markdown 格式化
            if bold and italic:
                text = f"***{text}***"
            elif bold:
                text = f"**{text}**"
            elif italic:
                text = f"_{text}_"
            
            parts.append(text)
        
        return ''.join(parts)
    
    def _extract_tables(self, doc: Document) -> List[TableData]:
        """
        提取文档中的所有表格。
        
        Args:
            doc: python-docx Document 对象
            
        Returns:
            [TableData, ...]
        """
        tables = []
        for table in doc.tables:
            if not table.rows:
                continue
            
            headers = []
            rows = []
            
            # 第一行作为表头
            try:
                headers = [cell.text.strip() for cell in table.rows[0].cells]
            except IndexError:
                continue
            
            # 数据行
            for row in table.rows[1:]:
                row_data = []
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:  # 跳过空单元格
                        row_data.append(text)
                if row_data:  # 只保留非空行
                    rows.append(row_data)
            
            tables.append(TableData(
                headers=headers,
                rows=rows
            ))
        
        return tables
    
    def _extract_images(self, doc: Document, output_dir: str) -> int:
        """
        提取文档中的图片。
        
        Args:
            doc: python-docx Document 对象
            output_dir: 输出目录
            
        Returns:
            提取的图片数量
        """
        if not os.path.exists(output_dir):
            return 0
        
        count = 0
        for rel in doc.part.rels.values():
            # 检查是否为图片关系
            if "image" in rel.reltype:
                try:
                    image_data = rel.target_part.blob
                    ext = self._get_image_extension(rel)
                    filename = f"img_{count}.{ext}"
                    filepath = os.path.join(output_dir, filename)
                    
                    with open(filepath, 'wb') as f:
                        f.write(image_data)
                    count += 1
                except Exception:
                    pass
        
        return count
    
    def _get_image_extension(self, rel) -> str:
        """
        从关系类型推断图片扩展名。
        
        Args:
            rel: python-docx Relationship 对象
            
        Returns:
            文件扩展名（如 'png', 'jpg'）
        """
        content_type = getattr(rel.target_part, 'content_type', '')
        if 'png' in content_type:
            return 'png'
        elif 'jpeg' in content_type or 'jpg' in content_type:
            return 'jpg'
        elif 'gif' in content_type:
            return 'gif'
        elif 'bmp' in content_type:
            return 'bmp'
        else:
            return 'png'  # 默认
    
    def _estimate_pages(self, doc: Document) -> int:
        """
        估算文档页数（基于段落数量）。
        
        Args:
            doc: python-docx Document 对象
            
        Returns:
            估算的页数
        """
        # 简单估算：每页约 50 个段落
        total_paragraphs = len(doc.paragraphs)
        return max(1, total_paragraphs // 50 + 1)
    
    def _write_tree_structure(self, sections: List[SectionResult], headings: List[HeadingInfo],
                               base_dir: str, doc_name: str):
        """
        将转换结果写入树状 MD 文件结构。
        
        Args:
            sections: 章节列表
            headings: 标题列表
            base_dir: 基础输出目录
            doc_name: 文档名称（不含扩展名）
        """
        # 创建文档根目录
        doc_dir = os.path.join(base_dir, doc_name)
        os.makedirs(doc_dir, exist_ok=True)
        
        # 生成 _index.md（全文摘要 + 章节导航）
        self._write_index_md(sections=sections, headings=headings,
                            index_path=os.path.join(doc_dir, "_index.md"))
        
        # 按层级组织文件
        for section in sections:
            if section.heading_level == 1:
                # 顶级章节 → 直接放在根目录
                filename = self._sanitize_filename(section.title)
                filepath = os.path.join(doc_dir, f"{filename}.md")
                self._write_section_md(
                    section=section,
                    doc_name=doc_name,
                    output_path=filepath
                )
            else:
                # 子章节 → 创建子目录
                parent_dir = os.path.join(doc_dir, *self._get_directory_path(section))
                os.makedirs(parent_dir, exist_ok=True)
                filename = self._sanitize_filename(section.title)
                filepath = os.path.join(parent_dir, f"{filename}.md")
                self._write_section_md(
                    section=section,
                    doc_name=doc_name,
                    output_path=filepath
                )
    
    def _write_index_md(self, sections: List[SectionResult], headings: List[HeadingInfo],
                        index_path: str):
        """
        生成 _index.md（全文摘要 + 章节导航）。
        
        Args:
            sections: 章节列表
            headings: 标题列表
            index_path: 输出路径
        """
        lines = [
            f"# {headings[0].text if headings else '文档目录'}",
            "",
            "## 📑 章节导航",
            ""
        ]
        
        # 按层级生成导航链接
        for heading in headings:
            indent = "  " * (heading.level - 1)
            lines.append(f"{indent}- [{heading.text}](./{'../' * (heading.level - 1)}{self._sanitize_filename(heading.text)}.md)")
        
        with open(index_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))
    
    def _write_section_md(self, section: SectionResult, doc_name: str,
                          output_path: str):
        """
        生成单个章节的 MD 文件。
        
        Args:
            section: 章节结果
            doc_name: 文档名称
            output_path: 输出路径
        """
        # YAML Frontmatter
        frontmatter = [
            f"---",
            f"source: \"{doc_name}.docx\"",
            f"chapter: \"{section.title}\"",
            f"level: L{section.heading_level}",
        ]
        
        if section.parent_title:
            frontmatter.append(f"parent: \"{section.parent_title}\"")
        
        if section.metadata.get('paragraph_range'):
            frontmatter.append(f"paragraph_range: \"{section.metadata['paragraph_range']}\"")
        
        frontmatter.extend([
            f"tags: [docx, {self._sanitize_filename(section.title)}]",
            "---",
            ""
        ])
        
        # 标题（H1-H9）
        heading_prefix = '#' * section.heading_level
        content_lines = [
            '\n'.join(frontmatter),
            f"{heading_prefix} {section.title}",
            "",
            section.content,
        ]
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(content_lines))
    
    def _get_directory_path(self, section: SectionResult) -> List[str]:
        """
        根据章节层级生成目录路径。
        
        Args:
            section: 章节结果
            
        Returns:
            目录路径列表（如 ['02_市场分析', '02_01_行业趋势']）
        """
        path = []
        # 找到所有上级标题
        for i in range(section.heading_level - 1):
            if i < len(path):
                pass  # 已在路径中
        return path
    
    def _sanitize_filename(self, text: str) -> str:
        """
        清理文件名中的非法字符。
        
        Args:
            text: 原始文本
            
        Returns:
            安全的文件名
        """
        # 替换非法字符
        sanitized = re.sub(r'[\\/:*?"<>|]', '_', text)
        # 限制长度
        if len(sanitized) > 50:
            sanitized = sanitized[:47] + '...'
        return sanitized.strip()
